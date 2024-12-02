import os
from anthropic import Anthropic, HUMAN_PROMPT, AI_PROMPT
from dotenv import load_dotenv
import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
from fpdf import FPDF  # PDF generation library
from docx import Document  # DOCX generation library

ANTHROPIC_API_KEY="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"

# Initialize the Anthropic client
client = anthropic = Anthropic(api_key=ANTHROPIC_API_KEY)

uploaded_cv_path = None  # Global variable to store the CV file path
def on_upload_cv_click():
    """Handles the CV upload and updates the UI."""
    global uploaded_cv_path
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf"), ("Word documents", "*.docx"), ("All files", "*.*")])
    if file_path:
        uploaded_cv_path = file_path  # Store the file path
        label_uploaded_cv.config(text=f"Uploaded CV: {os.path.basename(file_path)}")
        messagebox.showinfo("Success", f"CV uploaded successfully: {os.path.basename(file_path)}")
    else:
        label_uploaded_cv.config(text="No CV uploaded.")
        messagebox.showwarning("No File Selected", "You did not upload a CV.")

def generate_cover_letter(job_title, job_specifications, candidate_name, candidate_qualifications):
    """
    This function generates a formal cover letter using the Anthropic API based on the job title,
    job specifications, candidate name, and qualifications.
    """
    message = client.messages.create(
        model="claude-3-5-sonnet-20240620",
        max_tokens=1000,
        temperature=0,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"""
        Please generate a professional and formal cover letter for a job application. The candidate is applying for the following position:
        Job Title: {job_title}
        Job Specifications: {job_specifications}
        Candidate Name: {candidate_name}
        Candidate Qualifications: {candidate_qualifications}
        
        If the candidate has uploaded a CV, you may reference additional qualifications or information from the CV as appropriate.
        
        The cover letter should:
        - Be formal and professional
        - Clearly express the candidate's qualifications for the job
        - Be succinct and well-organized
        - Use proper formatting with paragraph breaks and clear structure
        
        The cover letter should be structured like this:
        1. Introduction: Brief introduction stating the job position and the candidate's interest
        2. Body: Discuss the candidate's qualifications, skills, and experiences
        3. Conclusion: Express enthusiasm for the position and willingness to discuss further
        
        Please avoid any unnecessary characters or formatting errors, and ensure the cover letter is clear and professional.
        """
        }
            ]
        }]
    )
    if message.content and isinstance(message.content, list):
        formatted_response = message.content[0].text.strip()
        formatted_response = formatted_response.replace("\\n", "\n")
        return formatted_response
    return None
    return message.content  # Return the generated cover letter

def on_generate_click():
    """
    This function is triggered when the 'Generate Cover Letter' button is clicked.
    It gathers input from the UI and generates the cover letter.
    """
    job_title = entry_job_title.get()
    job_specifications = entry_job_specifications.get()
    candidate_name = entry_candidate_name.get()
    candidate_qualifications = entry_candidate_qualifications.get()

    if not job_title or not job_specifications or not candidate_name or not candidate_qualifications:
        messagebox.showerror("Input Error", "All fields are required!")
        return

    # Inform the user that the letter is being generated
    label_result.config(text="Generating your cover letter...")

    # Generate the cover letter using the provided inputs
    cover_letter = generate_cover_letter(job_title, job_specifications, candidate_name, candidate_qualifications)
    
    # Display the generated cover letter in the text area
    text_cover_letter.delete(1.0, tk.END)  # Clear the text area
    text_cover_letter.insert(tk.END, cover_letter)

# Function to sanitize special characters in the text (e.g., bullet points)
def sanitize_text_for_pdf(text):
    """Sanitize the text by replacing problematic characters with safe alternatives."""
    return text.replace("\u2022", "*") 

# Function to save cover letter as PDF
def save_as_pdf(cover_letter):
    """Save the cover letter as a PDF file."""
    # Sanitize the cover letter text
    sanitized_cover_letter = sanitize_text_for_pdf(cover_letter)

    # Create the PDF object
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Add the cover letter text to the PDF
    pdf.multi_cell(0, 10, sanitized_cover_letter)

    # Ask the user for the filename and save the PDF file
    file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
    if file_path:
        pdf.output(file_path)
        messagebox.showinfo("Success", f"Cover letter saved as PDF: {file_path}")
    else:
        messagebox.showwarning("Save Cancelled", "You didn't choose a location to save the PDF.")

# Function to save cover letter as DOCX
def save_as_docx(cover_letter):
    """Save the cover letter as a DOCX file."""
    doc = Document()
    doc.add_heading("Cover Letter", 0)  # Add a title to the document
    doc.add_paragraph(cover_letter)  # Add the cover letter text

    # Ask the user for the filename and save the DOCX file
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")])
    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Success", f"Cover letter saved as DOCX: {file_path}")

# Callback function when 'Save as PDF' button is clicked
def on_save_pdf_click():
    """Trigger the save as PDF action."""
    cover_letter = text_cover_letter.get(1.0, tk.END)  # Get the generated cover letter text
    save_as_pdf(cover_letter)

# Callback function when 'Save as DOCX' button is clicked
def on_save_docx_click():
    """Trigger the save as DOCX action."""
    cover_letter = text_cover_letter.get(1.0, tk.END)  # Get the generated cover letter text
    save_as_docx(cover_letter)
    
# Style function to apply styles to widgets
def apply_styles(widget, style):
    if style == 'header':
        widget.config(font=("Arial", 16, 'bold'), fg="#333")
    elif style == 'input-field':
        widget.config(font=("Arial", 12), bd=1, relief="solid", width=50)
    elif style == 'button':
        widget.config(font=("Arial", 12), bg="#4CAF50", fg="white", relief="flat", pady=10)
    elif style == 'button_hover':
        widget.config(bg="#45a049")
    elif style == 'text-area':
        widget.config(font=("Arial", 12), bd=1, relief="solid", height=10, wrap=tk.WORD)

# Set up the UI
root = tk.Tk()
root.title("Cover Letter Generator")

# Label to display the uploaded CV file name
label_uploaded_cv = tk.Label(root, text="No CV uploaded.", font=("Arial", 10))
label_uploaded_cv.grid(row=9, column=0, columnspan=2, pady=5)
apply_styles(label_uploaded_cv, 'header')  # Apply style

# Create and place the input labels and fields
label_job_title = tk.Label(root, text="Job Title:")
label_job_title.grid(row=0, column=0, padx=10, pady=5)
apply_styles(label_job_title, 'header')
entry_job_title = tk.Entry(root)
entry_job_title.grid(row=0, column=1, padx=10, pady=5)
apply_styles(entry_job_title, 'input-field')

label_job_specifications = tk.Label(root, text="Job Specifications:")
label_job_specifications.grid(row=1, column=0, padx=10, pady=5)
apply_styles(label_job_specifications, 'header')
entry_job_specifications = tk.Entry(root)
entry_job_specifications.grid(row=1, column=1, padx=10, pady=5)
apply_styles(entry_job_specifications, 'input-field')

label_candidate_name = tk.Label(root, text="Candidate Name:")
label_candidate_name.grid(row=2, column=0, padx=10, pady=5)
apply_styles(label_candidate_name, 'header')
entry_candidate_name = tk.Entry(root)
entry_candidate_name.grid(row=2, column=1, padx=10, pady=5)
apply_styles(entry_candidate_name, 'input-field')

label_candidate_qualifications = tk.Label(root, text="Candidate Qualifications:")
label_candidate_qualifications.grid(row=3, column=0, padx=10, pady=5)
apply_styles(label_candidate_qualifications, 'header')
entry_candidate_qualifications = tk.Entry(root)
entry_candidate_qualifications.grid(row=3, column=1, padx=10, pady=5)
apply_styles(entry_candidate_qualifications, 'input-field')

# Button to upload CV
button_upload_cv = tk.Button(root, text="Upload CV", command=on_upload_cv_click)
button_upload_cv.grid(row=7, column=0, columnspan=2, pady=5)
apply_styles(button_upload_cv, 'button')

# Button to generate cover letter
button_generate = tk.Button(root, text="Generate Cover Letter", command=on_generate_click)
button_generate.grid(row=4, column=0, columnspan=2, pady=10)
apply_styles(button_generate, 'button')

# Button to save as PDF
button_save_pdf = tk.Button(root, text="Save as PDF", command=on_save_pdf_click)
button_save_pdf.grid(row=5, column=0, columnspan=2, pady=5)
apply_styles(button_save_pdf, 'button')

# Button to save as DOCX
button_save_docx = tk.Button(root, text="Save as DOCX", command=on_save_docx_click)
button_save_docx.grid(row=6, column=0, columnspan=2, pady=5)
apply_styles(button_save_docx, 'button')

# Label to display the result
label_result = tk.Label(root, text="", font=("Arial", 12))
label_result.grid(row=7, column=0, columnspan=2, pady=10)
apply_styles(label_result, 'header')

# Text area to display the generated cover letter
text_cover_letter = tk.Text(root, width=60, height=10, wrap=tk.WORD)
text_cover_letter.grid(row=8, column=0, columnspan=2, padx=10, pady=10)
apply_styles(text_cover_letter, 'text-area')

# Run the GUI
root.mainloop()